"""Extension-based dispatch for multi-format file CRUD.

The model sees four simple verbs (read/write/edit/delete). This module routes
each verb to the correct format handler based on the file extension.

Design decisions (Phase 1):
  - READS return clean text. The model reasons over text, so PDF/DOCX/XLSX/
    images are all flattened to a readable text representation.
  - WRITES create or overwrite. Tabular text (CSV-style) is parsed into real
    XLSX cells; newline-separated text becomes DOCX/PDF paragraphs.
  - EDITS are text find-and-replace, supported for text-like formats and DOCX.
    XLSX/PDF/images cannot be text-edited in place — regenerate with write.
  - Image OCR is read-only. PDF in-place edit is not supported (regenerate).

Heavy libraries are imported lazily inside each handler so a missing optional
dependency only affects that one format, never the whole tool registry.
"""

from pathlib import Path

# Extensions we treat as plain UTF-8 text.
TEXT_EXTS = {
    ".txt", ".md", ".markdown", ".rst", ".log", ".json", ".xml", ".yaml",
    ".yml", ".ini", ".cfg", ".toml", ".py", ".js", ".ts", ".tsx", ".jsx",
    ".html", ".css", ".sh", ".env", ".csv",  # csv is valid text too
}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".gif", ".webp"}


def _ext(path: str) -> str:
    return Path(path).suffix.lower()


# ----------------------------------------------------------------------------
# READ
# ----------------------------------------------------------------------------

def read_any(path: str) -> str:
    if not Path(path).exists():
        return f"Error: file not found: {path}"

    ext = _ext(path)
    try:
        if ext == ".pdf":
            return _read_pdf(path)
        if ext == ".docx":
            return _read_docx(path)
        if ext == ".xlsx":
            return _read_xlsx(path)
        if ext in IMAGE_EXTS:
            return _read_image_ocr(path)
        # Default: plain text (covers TEXT_EXTS and anything unknown).
        return _read_text(path)
    except Exception as e:
        return f"Error reading {ext or 'file'}: {e}"


def _read_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _words_to_text(words: list[dict]) -> str:
    """Rebuild clean text from pdfplumber words, inserting real spaces.

    pdfplumber's extract_text() often loses spaces between words in PDFs that
    don't encode explicit space glyphs (you get 'Starlette'sRequestobject').
    extract_words() detects word boundaries by horizontal gaps, so we group
    words into lines by their vertical position and join with spaces.
    """
    lines: dict[int, list[dict]] = {}
    for w in words:
        # Bucket by vertical position (~line). Round to merge same-line words.
        key = round(w["top"] / 3)
        lines.setdefault(key, []).append(w)

    out_lines = []
    for key in sorted(lines):
        row = sorted(lines[key], key=lambda w: w["x0"])
        out_lines.append(" ".join(w["text"] for w in row))
    return "\n".join(out_lines)


def _read_pdf(path: str) -> str:
    import pdfplumber

    out = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            words = page.extract_words(x_tolerance=1.5, y_tolerance=3)
            if words:
                text = _words_to_text(words)
            else:
                # Fallback for pages where word extraction finds nothing.
                text = page.extract_text() or ""
            out.append(f"--- Page {i} ---\n{text}")
    return "\n\n".join(out) if out else "(PDF contained no extractable text.)"


def _read_docx(path: str) -> str:
    import docx

    document = docx.Document(path)
    return "\n".join(p.text for p in document.paragraphs)


def _read_xlsx(path: str) -> str:
    import openpyxl

    wb = openpyxl.load_workbook(path, data_only=True)
    out = []
    for sheet in wb.worksheets:
        out.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            out.append(",".join(cells))
    return "\n".join(out)


def _read_image_ocr(path: str) -> str:
    try:
        import pytesseract
        from PIL import Image
    except ImportError as e:
        return f"Error: OCR requires pytesseract and pillow ({e})."

    try:
        text = pytesseract.image_to_string(Image.open(path))
    except pytesseract.TesseractNotFoundError:
        return (
            "Error: the Tesseract OCR engine is not installed. "
            "Install it with: sudo apt install tesseract-ocr"
        )
    return text.strip() or "(No text detected in image.)"


# ----------------------------------------------------------------------------
# WRITE (create / overwrite)
# ----------------------------------------------------------------------------

def write_any(path: str, content: str) -> str:
    ext = _ext(path)

    parent = Path(path).parent
    if str(parent):
        parent.mkdir(parents=True, exist_ok=True)

    try:
        if ext == ".pdf":
            return _write_pdf(path, content)
        if ext == ".docx":
            return _write_docx(path, content)
        if ext == ".xlsx":
            return _write_xlsx(path, content)
        if ext in IMAGE_EXTS:
            return f"Error: writing image files is not supported (OCR is read-only)."
        return _write_text(path, content)
    except Exception as e:
        return f"Error writing {ext or 'file'}: {e}"


def _write_text(path: str, content: str) -> str:
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return f"File written: {path}"


def _write_pdf(path: str, content: str) -> str:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    doc = SimpleDocTemplate(path, pagesize=letter)
    styles = getSampleStyleSheet()
    flow = []
    for line in content.split("\n"):
        if line.strip():
            flow.append(Paragraph(line, styles["Normal"]))
        flow.append(Spacer(1, 6))
    doc.build(flow)
    return f"PDF created: {path}"


def _write_docx(path: str, content: str) -> str:
    import docx

    document = docx.Document()
    for line in content.split("\n"):
        document.add_paragraph(line)
    document.save(path)
    return f"DOCX created: {path}"


def _write_xlsx(path: str, content: str) -> str:
    import csv
    import io
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    reader = csv.reader(io.StringIO(content))
    for row in reader:
        ws.append(row)
    wb.save(path)
    return f"XLSX created: {path}"


# ----------------------------------------------------------------------------
# EDIT (text find-and-replace where the format allows it)
# ----------------------------------------------------------------------------

def edit_any(path: str, old_str: str, new_str: str) -> str:
    if not Path(path).exists():
        return f"Error: file not found: {path}"

    ext = _ext(path)
    try:
        if ext == ".docx":
            return _edit_docx(path, old_str, new_str)
        if ext in {".pdf", ".xlsx"} or ext in IMAGE_EXTS:
            return (
                f"Error: in-place edit is not supported for {ext} files. "
                f"Read it, change the text, and use write_file to regenerate."
            )
        return _edit_text(path, old_str, new_str)
    except Exception as e:
        return f"Error editing {ext or 'file'}: {e}"


def _edit_text(path: str, old_str: str, new_str: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    if old_str not in content:
        return f"Error: text not found in {path}"
    with open(path, "w", encoding="utf-8") as f:
        f.write(content.replace(old_str, new_str, 1))
    return f"File edited: {path}"


def _edit_docx(path: str, old_str: str, new_str: str) -> str:
    import docx

    document = docx.Document(path)
    replaced = False
    for para in document.paragraphs:
        if old_str in para.text:
            # Rewrite the paragraph text. Note: run-level formatting on the
            # edited paragraph may be flattened. Good enough for plain docs.
            new_text = para.text.replace(old_str, new_str, 1)
            for run in list(para.runs):
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
            replaced = True
            break
    if not replaced:
        return f"Error: text not found in {path}"
    document.save(path)
    return f"DOCX edited: {path}"
